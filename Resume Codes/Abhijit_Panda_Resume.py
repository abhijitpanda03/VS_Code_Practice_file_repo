from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------- Helper: Add Clickable Hyperlink ----------
def add_hyperlink(paragraph, text, url):
    """
    Add a clickable hyperlink within a paragraph.
    Only displays 'text' but links to 'url'.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # Blue color
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# ---------- Helper: Text Utilities ----------
def add_heading(doc, text, size=12, bold=True, color=RGBColor(0, 102, 204)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run("──────────────────────────────────────────────")
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# ---------- Create the Resume ----------
doc = Document()

# Global font style
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

# ---------- Header ----------
header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

name = header.add_run("ABHIJIT PANDA\n")
name.bold = True
name.font.size = Pt(16)

contact_line = header.add_run("Data Analyst - A | Pune, India | +91-8114957735 | abhijeetpanda2000@gmail.com\n")
contact_line.font.size = Pt(10)

# Add clickable LinkedIn link
linkedin_para = doc.add_paragraph()
linkedin_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_hyperlink(linkedin_para, "Abhijit|LinkedIn", "https://www.linkedin.com/in/abhijit-panda-68395618b/")

add_separator(doc)

# ---------- Professional Summary ----------
add_heading(doc, "PROFESSIONAL SUMMARY")
summary = (
    "Results-oriented Azure Data Engineer with 4+ years of experience designing and delivering end-to-end "
    "data migration and ETL solutions using Azure Data Factory, Databricks, and Synapse Analytics. "
    "Demonstrated expertise in modernizing on-premises systems to cloud-native architectures, optimizing data "
    "pipelines for performance and scalability, and implementing automated validation frameworks to ensure high "
    "data quality. Skilled at translating complex business requirements into robust, high-performing data workflows "
    "that power analytics and strategic decision-making."
)
doc.add_paragraph(summary)
add_separator(doc)

# ---------- Experience ----------
add_heading(doc, "EXPERIENCE")

# CAP
p = doc.add_paragraph()
p.add_run("CAPGEMINI | Data Analyst – A | Dec 2024 - Present").bold = True
doc.add_paragraph("Project – Azure Data Migration", style="List Bullet")

add_bullet(doc, "Led migration of on-premises systems to Azure, designing and managing ADF and Synapse pipelines to move and transform enterprise datasets into Delta Lake.")
add_bullet(doc, "Developed PySpark and SparkSQL-based transformations to handle complex business logic across multiple data sources and formats.")
add_bullet(doc, "Optimized ADF and Synapse performance, reducing pipeline runtime by 30% and improving data reliability.")
add_bullet(doc, "Automated data validation and reconciliation frameworks, increasing data accuracy by 90% and minimizing manual intervention.")
add_bullet(doc, "Delivered curated datasets for Power BI analytics, supporting financial planning, forecasting, and operational reporting.")

# ACN
p = doc.add_paragraph()
p.add_run("ACCENTURE | Data Engineering Analyst | Aug 2021 – Dec 2024").bold = True
doc.add_paragraph("Project – Finance Data and Experience", style="List Bullet")

add_bullet(doc, "Built and maintained ADF pipelines to migrate and transform sales data marts from SQL Server to Azure Data Lake (ADLS).")
add_bullet(doc, "Created Synapse notebooks for large-scale data transformations, aggregations, and business rule validations.")
add_bullet(doc, "Implemented audit and reconciliation mechanisms to ensure accuracy between raw and transformed data layers.")
add_bullet(doc, "Supported forecasting and budgeting reports through validated, aggregated data outputs for Power BI.")
add_bullet(doc, "Participated in production support and Hypercare activities, troubleshooting and resolving pipeline issues efficiently.")

add_separator(doc)

# ---------- Skills ----------
add_heading(doc, "SKILLS")
skills = (
    "• Azure Data Factory • Databricks • Synapse Analytics • Azure Data Lake\n"
    "• PySpark • Python • SQL • Git/Bitbucket • Azure DevOps (basic)\n"
    "• ETL Development • Data Migration • Performance Optimization"
)
doc.add_paragraph(skills)
add_separator(doc)

# ---------- Certifications ----------
add_heading(doc, "CERTIFICATIONS")
certs = (
    "• Microsoft Certified: Azure Data Engineer Associate (DP-203)\n"
    "• Microsoft Certified: Azure Fundamentals (AZ-900)\n"
    "• Databricks Certified Data Engineer Associate\n"
    "• Python for Everybody (Coursera)"
)
doc.add_paragraph(certs)
add_separator(doc)

# ---------- Awards & Achievements ----------
add_heading(doc, "AWARDS & ACHIEVEMENTS")
add_bullet(doc, "Client Value Creation Award – Recognized for delivering high-quality data solutions that improved business insights.")
add_bullet(doc, "Integrity Award – For maintaining high standards of data governance and security.")
add_bullet(doc, "Inspiring Innovator – Honored for proposing automation initiatives that reduced manual validation by 90%.")
add_bullet(doc, "Accenture Achievement Recognition – For consistent performance and team collaboration in key Azure data projects.")

# ---------- Save ----------
output_path = "Abhijit_Panda_Resume_ModernV1.docx"
doc.save(output_path)

print(f"✅ Resume created successfully! Open this file:\n{output_path}")
